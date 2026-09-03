#!/usr/bin/env python3
"""
Generates Harika's complete beginner guide DOCX for Project SUTRA.
Includes role overview, GitHub tutorial, Gate verification instructions.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

IMG_DIR = "/home/nikhil/.gemini/antigravity-cli/brain/b16a6917-8627-44c3-a432-3640f1d4ff72"
OUT_PATH = "/home/nikhil/Desktop/Project SUTRA/docs/HARIKA_ROLE_GUIDE.docx"

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def heading(doc, text, level=1, color=RGBColor(0x0A, 0x2A, 0x5E)):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p

def body(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    return p

def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    # shade background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p

def divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0A2A5E')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(10)

def add_image(doc, path, caption, width=Inches(5.5)):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.italic = True
        cap.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    else:
        doc.add_paragraph(f"[Image missing: {path}]")

def note_box(doc, text, color="FFF3CD"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, color)
    cell.width = Inches(5.5)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    doc.add_paragraph()

# ─── Build Document ────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ══════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run("🛸  PROJECT SUTRA")
r.font.size = Pt(28)
r.font.bold = True
r.font.color.rgb = RGBColor(0x0A, 0x2A, 0x5E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Subsystem E — Harika's Complete Role & GitHub Guide")
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x1A, 0x7A, 0x5A)
r2.font.bold = True

doc.add_paragraph()
note_box(doc, "📌  This document is written for absolute beginners. No coding experience required. Just follow the steps exactly as shown and you will be fine!", "E8F4FD")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# SECTION 1 — WHAT IS PROJECT SUTRA?
# ══════════════════════════════════════════════════════════════════════
heading(doc, "1. What is Project SUTRA?", 1)
body(doc, "Project SUTRA (Swarm Unified Tactical Reconnaissance Architecture) is a team project building an autonomous drone swarm that can search for survivors in disaster zones — even in areas with no GPS or internet signal.")
body(doc, "The project has 5 subsystems (teams), each responsible for one part of the drone system:")
doc.add_paragraph()

# Team table
tbl = doc.add_table(rows=6, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Subsystem", "What it Does", "Lead"]
colors  = ["0A2A5E", "1E5C3A", "7A1E1E", "4A1E7A", "1A5A7A", "1A3D7A"]
row_data = [
    ("A — GNC",           "Makes drones fly automatically and avoid collisions",                       "Rohith Kumar"),
    ("B — Comms",         "Handles drone-to-drone communication & data compression",                   "Nikhil"),
    ("C — AI Perception", "AI camera: detects survivors, threats using YOLOv8",                        "Vedanth Sai Ram"),
    ("D — GCS Dashboard", "3D map on a computer showing all drones live",                              "Siva Kesava"),
    ("E — Docs & Audit",  "Writes all documents, verifies the whole system works — YOUR ROLE",         "Harika ⭐"),
]
header_cells = tbl.rows[0].cells
for i, h in enumerate(headers):
    header_cells[i].text = h
    set_cell_bg(header_cells[i], "0A2A5E")
    header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    header_cells[i].paragraphs[0].runs[0].font.bold = True

for ri, (name, desc, lead) in enumerate(row_data):
    row = tbl.rows[ri + 1]
    row.cells[0].text = name
    row.cells[1].text = desc
    row.cells[2].text = lead
    bg = "FFF9E6" if ri == 4 else "F7FBFF"
    for c in row.cells:
        set_cell_bg(c, bg)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════
# SECTION 2 — HARIKA'S ROLE OVERVIEW
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading(doc, "2. Your Role — Harika (Subsystem E Lead)", 1)
body(doc, "You are the Quality Gate Guardian and Documentation Lead of the entire project. Think of yourself as the project's 'Inspector' and 'Reporter'.")
body(doc, "You do NOT need to write any drone code. Your job has 3 clear parts:")

doc.add_paragraph()
add_image(doc,
    f"{IMG_DIR}/harika_role_overview_1785399277445.jpg",
    "Figure 1 — Harika sits at the centre: all 4 subsystems feed into her for verification",
    width=Inches(5.8))

doc.add_paragraph()
heading(doc, "Your 3 Core Jobs", 2)

heading(doc, "Job 1 — Run the Verification Test Suite (Gate Auditor)", 3, RGBColor(0x0A, 0x5E, 0x2A))
body(doc, "You run a single Python script that automatically checks whether all 5 subsystems are working correctly. The script checks 6 'Gates' (pass/fail tests). ALL 6 must pass before the team can submit.")
code_block(doc, "python3 scripts/SUTRA_48Hr_Hackathon_Master_Suite.py")
body(doc, "If any gate fails, you note which one failed and tell the responsible teammate so they can fix it.")

doc.add_paragraph()
add_image(doc,
    f"{IMG_DIR}/gate_verification_checklist_1785399323501.jpg",
    "Figure 2 — The 6 Verification Gates you must confirm ALL pass before submission",
    width=Inches(5.5))

doc.add_paragraph()

# Gate table
heading(doc, "Gate Reference Table", 3, RGBColor(0x0A, 0x2A, 0x5E))
gtbl = doc.add_table(rows=7, cols=4)
gtbl.style = 'Table Grid'
gtbl.alignment = WD_TABLE_ALIGNMENT.CENTER
gh = ["Gate", "What it Checks", "Must Show", "Which Teammate Fixes If Failed"]
for i, h in enumerate(gh):
    c = gtbl.rows[0].cells[i]
    c.text = h
    set_cell_bg(c, "0A2A5E")
    c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    c.paragraphs[0].runs[0].font.bold = True
gates = [
    ("G1","Physics simulation speed","RTF ≥ 0.98","Nikhil (B)"),
    ("G2","Drone-to-drone comms latency","< 8 ms, Failover < 150 ms","Nikhil (B)"),
    ("G3","AI survivor detection accuracy","mAP@0.5 ≥ 94%","Vedanth (C)"),
    ("G4","GPS target location accuracy","Error < 0.8 m","Vedanth (C)"),
    ("G5","Drone collision avoidance gap","Buffer > 2.8 m","Rohith (A)"),
    ("G6","Ground Control dashboard framerate","60 FPS locked","Siva Kesava (D)"),
]
gate_colors = ["E8F5E9","E3F2FD","FFF3E0","F3E5F5","E0F7FA","FCE4EC"]
for ri, (g, desc, thresh, fix) in enumerate(gates):
    row = gtbl.rows[ri + 1]
    row.cells[0].text = g
    row.cells[1].text = desc
    row.cells[2].text = thresh
    row.cells[3].text = fix
    for c in row.cells:
        set_cell_bg(c, gate_colors[ri])

doc.add_paragraph()
divider(doc)

# ──────────────────────────────────────────────────────────────────────
heading(doc, "Job 2 — Update & Audit the Subsystem DOCS.md Files", 3, RGBColor(0x0A, 0x5E, 0x2A))
body(doc, "Each subsystem has a dedicated documentation file (DOCS.md). These files record performance numbers like speed, accuracy, and memory usage. Your job is to:")
bullet(doc, "Read the DOCS.md file for each subsystem.")
bullet(doc, "Check that the numbers in the DOCS.md file match what the verification script actually measured.")
bullet(doc, "If they don't match, update the DOCS.md with the correct numbers from the test output.")
body(doc, "The 5 DOCS.md files you audit are:")
bullet(doc, "sutra_ws/src/sutra_gnc/DOCS.md  (Rohith's GNC system)", bold_prefix="A:")
bullet(doc, "sutra_ws/src/sutra_comms/DOCS.md  (Nikhil's comms system)", bold_prefix="B:")
bullet(doc, "sutra_ws/src/sutra_perception/DOCS.md  (Vedanth's AI)", bold_prefix="C:")
bullet(doc, "sutra_ws/src/sutra_gcs/DOCS.md  (Siva's dashboard)", bold_prefix="D:")
bullet(doc, "sutra_ws/src/sutra_sim/DOCS.md  (Simulation world)", bold_prefix="E:")

doc.add_paragraph()
note_box(doc, "💡 TIP: You don't need to understand the technical content deeply. Just compare the numbers. If the script says mAP = 94.8% but the DOCS.md says 90%, update the DOCS.md to say 94.8%.", "E8F4FD")

divider(doc)

# ──────────────────────────────────────────────────────────────────────
heading(doc, "Job 3 — Write the Project Whitepaper & Presentation Script", 3, RGBColor(0x0A, 0x5E, 0x2A))
body(doc, "You also write and maintain the team's formal documentation stored in the docs/ folder:")
bullet(doc, "System Architecture Whitepaper — explains how all 5 subsystems work together")
bullet(doc, "Hackathon Presentation Script — what each teammate will say during the demo")
bullet(doc, "Flight Logs — a record of every test run with results")
bullet(doc, "Roadmaps — track which features are done / in progress / pending")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════
# SECTION 3 — GITHUB BEGINNER TUTORIAL
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading(doc, "3. GitHub Tutorial — Step-by-Step for Beginners", 1)
body(doc, "GitHub is where the entire team's code and documents are stored online. Think of it like Google Drive, but for code. You don't need to touch code — you only need to:")
bullet(doc, "Download the latest files from GitHub to your laptop")
bullet(doc, "Make your changes (edit DOCS.md files, add reports)")
bullet(doc, "Upload your changes back to GitHub")

heading(doc, "3.1 — One-Time Setup", 2)
heading(doc, "Step 1: Install Git", 3, RGBColor(0x4A, 0x1E, 0x7A))
body(doc, "Git is the program that lets you connect your laptop to GitHub.")
bullet(doc, "Go to: https://git-scm.com/downloads")
bullet(doc, "Click 'Download for Windows' (or Mac)")
bullet(doc, "Install it with all default settings")
bullet(doc, "Open 'Git Bash' (Windows) or Terminal (Mac/Linux)")
code_block(doc, "git --version\n# Should show something like: git version 2.x.x")

heading(doc, "Step 2: Configure Git with Your Name", 3, RGBColor(0x4A, 0x1E, 0x7A))
body(doc, "This tells Git who you are so your changes show your name. Run these two lines in Git Bash:")
code_block(doc, 'git config --global user.name "Harika"\ngit config --global user.email "your-email@example.com"')

heading(doc, "Step 3: Clone the Repository (Download the Project)", 3, RGBColor(0x4A, 0x1E, 0x7A))
body(doc, "This downloads the entire project to your laptop. You only do this ONCE.")
code_block(doc, "git clone https://github.com/nikhil49023/SUTRA.git\ncd SUTRA")
body(doc, "After this, you will see a folder called SUTRA on your laptop with all the project files.")

doc.add_paragraph()
note_box(doc, "⚠️  IMPORTANT: You only clone once. Every time after that, just open the SUTRA folder and run git pull to get the latest updates.", "FFF3CD")

divider(doc)
heading(doc, "3.2 — Daily Workflow (Every Time You Work)", 2)

heading(doc, "Step 4: Open Your Project Folder in Terminal", 3, RGBColor(0x4A, 0x1E, 0x7A))
body(doc, "On Windows: Right-click inside the SUTRA folder → 'Open Git Bash here'")
body(doc, "On Mac/Linux: Open Terminal and type:")
code_block(doc, "cd ~/Desktop/SUTRA   # (or wherever you saved it)")

heading(doc, "Step 5: Switch to YOUR Branch", 3, RGBColor(0x4A, 0x1E, 0x7A))
body(doc, "A 'branch' is like your personal copy of the project where you make changes. Your branch is called feature/subsystem-e-docs. Always switch to it before working.")
code_block(doc, "git checkout feature/subsystem-e-docs")
body(doc, "You should see: Switched to branch 'feature/subsystem-e-docs'")

heading(doc, "Step 6: Pull the Latest Changes from the Team", 3, RGBColor(0x4A, 0x1E, 0x7A))
body(doc, "This downloads any updates your teammates made since you last worked. ALWAYS do this before making any changes.")
code_block(doc, "git fetch origin dev\ngit merge origin/dev --no-edit")
body(doc, "If it says 'Already up to date.' — great, you already have the latest version!")

note_box(doc, "📌 RULE #0: This fetch + merge step is MANDATORY at the start of every session. Never skip it!", "FFE8E8")

heading(doc, "Step 7: Run the Verification Suite", 3, RGBColor(0x4A, 0x1E, 0x7A))
body(doc, "Now run the master verification script and watch the output:")
code_block(doc, "python3 scripts/SUTRA_48Hr_Hackathon_Master_Suite.py")
body(doc, "You will see 6 lines of output. All 6 must end with ✓. If any show an error, note the Gate number and message, and tell the responsible teammate.")

heading(doc, "Step 8: Make Your Changes", 3, RGBColor(0x4A, 0x1E, 0x7A))
body(doc, "Now open and edit whichever DOCS.md or docs/ file you need to update. You can use any text editor — Notepad, VS Code, or even Notepad++.")
body(doc, "For example, to update Rohith's DOCS.md:")
code_block(doc, "# Open this file in your text editor:\nsutra_ws/src/sutra_gnc/DOCS.md")

heading(doc, "Step 9: Save Your Changes to GitHub (Commit & Push)", 3, RGBColor(0x4A, 0x1E, 0x7A))
body(doc, "After editing files, run these 3 commands to upload your work:")
code_block(doc, '# Step 9a: Stage all your changes\ngit add .\n\n# Step 9b: Write a commit message describing what you did\ngit commit -m "audit(gates): updated DOCS.md files with G1-G6 verified results"\n\n# Step 9c: Push (upload) to GitHub\ngit push origin feature/subsystem-e-docs')
body(doc, "After this, your changes are live on GitHub for the whole team to see!")

divider(doc)
heading(doc, "3.3 — Quick Reference Card", 2)

ref_tbl = doc.add_table(rows=7, cols=2)
ref_tbl.style = 'Table Grid'
ref_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
rh = ref_tbl.rows[0].cells
rh[0].text = "What You Want to Do"
rh[1].text = "Command to Run"
set_cell_bg(rh[0], "1A3D7A"); set_cell_bg(rh[1], "1A3D7A")
for c in rh:
    c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    c.paragraphs[0].runs[0].font.bold = True

ref_rows = [
    ("Download project (first time only)",   "git clone https://github.com/nikhil49023/SUTRA.git"),
    ("Switch to your branch",                "git checkout feature/subsystem-e-docs"),
    ("Get latest team changes",              "git fetch origin dev && git merge origin/dev --no-edit"),
    ("Run verification gates",               "python3 scripts/SUTRA_48Hr_Hackathon_Master_Suite.py"),
    ("Stage changes",                        "git add ."),
    ("Save + upload your changes",           'git commit -m "your message" && git push origin feature/subsystem-e-docs'),
]
for ri, (action, cmd) in enumerate(ref_rows):
    row = ref_tbl.rows[ri + 1]
    row.cells[0].text = action
    row.cells[1].text = cmd
    row.cells[1].paragraphs[0].runs[0].font.name = 'Courier New'
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
    bg = "F7FBFF" if ri % 2 == 0 else "FFFFFF"
    for c in row.cells:
        set_cell_bg(c, bg)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════
# SECTION 4 — COMMON PROBLEMS & SOLUTIONS
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading(doc, "4. Common Problems & Solutions", 1)

problems = [
    ("❌ I ran git checkout and got an error",
     "You might be on the wrong directory. Make sure you cd into the SUTRA folder first:\n  cd ~/Desktop/SUTRA\n  git checkout feature/subsystem-e-docs"),
    ("❌ The merge says 'CONFLICT'",
     "Don't panic. Just message Nikhil on WhatsApp/Teams with a screenshot. He will fix it for you since he has full access."),
    ("❌ python3 not found when running the script",
     "You need to install Python 3. Go to https://python.org → Download → Install. Make sure to check 'Add Python to PATH' during install."),
    ("❌ A gate fails (shows AssertionError)",
     "Note the Gate number (e.g. G3) and the error message. Send it to the responsible teammate (check the Gate table in Section 2). Do NOT try to fix it yourself."),
    ("❌ git push says 'rejected'",
     "Run git pull origin feature/subsystem-e-docs first, then push again."),
]

for prob, sol in problems:
    note_box(doc, f"{prob}\n\n✅ Solution: {sol}", "FFF0F0")

# ══════════════════════════════════════════════════════════════════════
# SECTION 5 — DAILY CHECKLIST
# ══════════════════════════════════════════════════════════════════════
heading(doc, "5. Harika's Daily Checklist", 1)
body(doc, "Print this page or save it on your phone. Every time you sit down to work, follow this list in order:")
doc.add_paragraph()

checklist = [
    ("☐ 1", "Open Git Bash inside the SUTRA folder"),
    ("☐ 2", "Switch to your branch:  git checkout feature/subsystem-e-docs"),
    ("☐ 3", "Pull latest changes:  git fetch origin dev && git merge origin/dev"),
    ("☐ 4", "Run verification:  python3 scripts/SUTRA_48Hr_Hackathon_Master_Suite.py"),
    ("☐ 5", "Check all 6 Gates show ✓  (if any fail, notify the responsible teammate)"),
    ("☐ 6", "Update any DOCS.md files with numbers from the test output"),
    ("☐ 7", "Stage, commit, and push:  git add . → git commit → git push"),
    ("☐ 8", "Notify the team on chat: 'Gate audit done — all 6 passed ✅'"),
]
ctl = doc.add_table(rows=len(checklist), cols=2)
ctl.style = 'Table Grid'
ctl.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, (cb, text) in enumerate(checklist):
    row = ctl.rows[ri]
    row.cells[0].text = cb
    row.cells[0].width = Inches(0.5)
    row.cells[1].text = text
    bg = "F0FFF4" if ri % 2 == 0 else "FFFFFF"
    for c in row.cells:
        set_cell_bg(c, bg)

doc.add_paragraph()
note_box(doc, "🏆 Remember: You are the last line of defence before the hackathon demo. If all 6 gates pass and the docs are accurate, the team is READY TO WIN!", "E8F5E9")

# ══════════════════════════════════════════════════════════════════════
# FOOTER NOTE
# ══════════════════════════════════════════════════════════════════════
divider(doc)
foot = doc.add_paragraph("Project SUTRA — Subsystem E Guide | Document maintained by Nikhil (Tech Lead) | For questions: contact Nikhil")
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot.runs[0].font.size = Pt(9)
foot.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.save(OUT_PATH)
print(f"✅ Document saved to: {OUT_PATH}")
